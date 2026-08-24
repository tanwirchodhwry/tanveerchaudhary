import streamlit as st
import fitz  # PyMuPDF
import ollama
from pydantic import BaseModel, Field
from typing import List, Optional
import json
import pandas as pd
from docx import Document
from io import BytesIO

# --- 1. Define Data Structures for AI ---
class Experience(BaseModel):
    company: str
    job_title: str
    years_of_experience: float
    skills_used: List[str]

class ResumeData(BaseModel):
    name: str 
    email: str
    phone: Optional[str]
    total_years_experience: float
    top_skills: List[str]
    work_history: List[Experience]

class MatchScore(BaseModel):
    score: int
    reason: str

# --- 2. Processing Pipeline Functions ---
def extract_text(uploaded_file):
    doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
    text = ""
    for page in doc:
        text += page.get_text()
    return text

def parse_with_ollama(text):
    try:
        response = ollama.chat(
            model='llama3.2:1b',
            messages=[
                {"role": "system", "content": "You are an expert HR assistant. Extract fields accurately."},
                {"role": "user", "content": text}
            ],
            format=ResumeData.model_json_schema(),
        )
        return json.loads(response['message']['content'])
    except Exception:
        return {}

def calculate_match_score(candidate_skills, job_description):
    if not job_description.strip():
        return 0, "No Job Description provided"
    try:
        prompt = f"Job Description:\n{job_description}\n\nCandidate Skills:\n{candidate_skills}"
        response = ollama.chat(
            model='llama3.2:1b',
            messages=[
                {"role": "system", "content": "Rate candidate fit from 0 to 100 based on job criteria."},
                {"role": "user", "content": prompt}
            ],
            format=MatchScore.model_json_schema(),
        )
        result = json.loads(response['message']['content'])
        return result.get('score', 0), result.get('reason', 'Processed successfully')
    except Exception:
        return 0, "Error calculating score"

# --- 3. Word Document Generator ---
def generate_word_document(dataframe):
    doc = Document()
    doc.add_heading("AI Talent Pipeline Screening Report", level=0)
    doc.add_paragraph("Automated candidate shortlists compiled by the local AI screening engine.")
    doc.add_paragraph("="*60)
    
    for _, row in dataframe.iterrows():
        doc.add_heading(f"👤 {row['Candidate Name']}", level=1)
        p = doc.add_paragraph()
        p.add_run("• Match Grade: ").bold = True
        p.add_run(f"{row['Match Score (%)']}%\n")
        p.add_run("• Relevant Experience: ").bold = True
        p.add_run(f"{row['Experience (Years)']} Years\n")
        p.add_run("• Contact Registry: ").bold = True
        p.add_run(f"{row['Email']}\n")
        
        doc.add_heading("Extracted Core Competencies:", level=2)
        doc.add_paragraph(str(row['Top Skills Extracted']))
        
        doc.add_heading("AI Fit Evaluation Summary:", level=2)
        doc.add_paragraph(str(row['AI Evaluation Summary']))
        doc.add_paragraph("-" * 40)
        
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# --- 4. User Interface View ---
st.set_page_config(page_title="AI Resume Pipeline", layout="wide")
st.title("🎯 Enterprise AI Resume Parser & Ranking Dashboard")

# Interface Sidebar Configuration
st.sidebar.header("🔧 System Configurations")
job_desc = st.sidebar.text_area("📋 Paste target Job Description here:", height=200)
min_experience = st.sidebar.slider("⏳ Minimum Required Experience (Years)", 0, 15, 0)

# Main File Upload Layout
uploaded_files = st.file_uploader("Choose candidate resumes (PDF format)", type=["pdf"], accept_multiple_files=True)

if uploaded_files:
    if st.button("🚀 Execute Ingestion & Parsing Pipeline", type="primary"):
        all_candidates_data = []
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        for index, file in enumerate(uploaded_files):
            status_text.text(f"Processing ({index + 1}/{len(uploaded_files)}): {file.name}...")
            
            raw_text = extract_text(file)
            parsed_json = parse_with_ollama(raw_text)
            
            skills_string = ", ".join(parsed_json.get('top_skills', []))
            exp_years = float(parsed_json.get('total_years_experience', 0))
            score, reason = calculate_match_score(skills_string, job_desc)
            
            candidate_row = {
                "Filename": file.name,
                "Candidate Name": parsed_json.get('name', 'N/A'),
                "Email": parsed_json.get('email', 'N/A'),
                "Experience (Years)": exp_years,
                "Top Skills Extracted": skills_string,
                "Match Score (%)": score,
                "AI Evaluation Summary": reason
            }
            all_candidates_data.append(candidate_row)
            progress_bar.progress((index + 1) / len(uploaded_files))
            
        status_text.text("🎉 Batch execution complete!")
        st.session_state['master_df'] = pd.DataFrame(all_candidates_data)

# Live Table Presentation
if 'master_df' in st.session_state and not st.session_state['master_df'].empty:
    df_working = st.session_state['master_df'].copy()
    filtered_df = df_working[df_working["Experience (Years)"] >= min_experience]
    filtered_df = filtered_df.sort_values(by="Match Score (%)", ascending=False)
    
    st.write("---")
    st.subheader(f"📋 Screened Candidate Shortlist ({len(filtered_df)} profiles matching criteria)")
    st.dataframe(filtered_df, use_container_width=True, hide_index=True)
    
    col_csv, col_docx = st.columns(2)
    with col_csv:
        st.download_button(
            label="📥 Download Data Matrix Sheet (CSV)",
            data=filtered_df.to_csv(index=False).encode('utf-8'),
            file_name="screened_candidates.csv",
            mime="text/csv",
        )
    with col_docx:
        st.download_button(
            label="📄 Download Executive Document (Word)",
            data=generate_word_document(filtered_df),
            file_name="screening_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
